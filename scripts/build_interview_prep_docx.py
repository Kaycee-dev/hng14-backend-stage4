"""
Build interview_prep.docx from interview_prep.md.

Style: navy/blue colour scheme matching the DevOps stage4 prep doc format.
  Heading colour: #1A1A2E
  Accent colour:  #1B6CA8
  Code blocks: monospace, grey background
  Q&A pairs: question in accent blue with left border, answer indented
  Tables: navy header rows, alternating row shading
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import lxml.etree as etree

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT  = os.path.dirname(SCRIPT_DIR)
MD_PATH    = os.path.join(REPO_ROOT, 'docs', 'backend-stage4-control-plane', 'interview_prep.md')
DOCX_PATH  = os.path.join(REPO_ROOT, 'docs', 'backend-stage4-control-plane', 'interview_prep.docx')

# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------
NAVY        = RGBColor(0x1A, 0x1A, 0x2E)   # heading colour
ACCENT      = RGBColor(0x1B, 0x6C, 0xA8)   # accent blue
LIGHT_BLUE  = RGBColor(0xD6, 0xE4, 0xF0)   # table row alt shading
CODE_BG     = RGBColor(0xF2, 0xF2, 0xF2)   # code block background
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY   = RGBColor(0x44, 0x44, 0x44)

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def rgb_hex(colour: RGBColor) -> str:
    return '{:02X}{:02X}{:02X}'.format(colour[0], colour[1], colour[2])


def set_run_colour(run, colour: RGBColor):
    rpr = run._r.get_or_add_rPr()
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), rgb_hex(colour))
    rpr.append(clr)


def set_para_shading(para, fill: RGBColor):
    ppr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(fill))
    ppr.append(shd)


def set_cell_shading(cell, fill: RGBColor):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(fill))
    tcpr.append(shd)


def add_para_border_left(para, colour: RGBColor, space: int = 6, sz: int = 12):
    ppr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    hex_c = rgb_hex(colour)
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), str(space))
    left.set(qn('w:color'), hex_c)
    pBdr.append(left)
    ppr.append(pBdr)


def set_cell_borders(table):
    """Apply thin borders to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcpr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), rgb_hex(NAVY))
                tcBorders.append(border)
            tcpr.append(tcBorders)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level):
    para = doc.add_heading(text, level=level)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.color.rgb = NAVY
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)
    return para


def add_code_block(doc, text):
    para = doc.add_paragraph()
    set_para_shading(para, CODE_BG)
    ppr = para._p.get_or_add_pPr()
    # small indent
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ppr.append(ind)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_GREY
    return para


def add_question_para(doc, text):
    para = doc.add_paragraph()
    add_para_border_left(para, ACCENT, space=8, sz=16)
    set_para_shading(para, RGBColor(0xEB, 0xF4, 0xFB))
    run = para.add_run(text)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(10)
    return para


def add_answer_para(doc, text):
    para = doc.add_paragraph()
    ppr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ppr.append(ind)
    run = para.add_run(text)
    run.font.size = Pt(10)
    return para


def add_body_para(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    return para


def add_table_from_rows(doc, header_row, data_rows):
    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = 'Table Grid'

    # Header row
    hrow = table.rows[0]
    for i, cell_text in enumerate(header_row):
        cell = hrow.cells[i]
        set_cell_shading(cell, NAVY)
        run = cell.paragraphs[0].add_run(cell_text)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        row = table.rows[r_idx + 1]
        fill = LIGHT_BLUE if r_idx % 2 == 0 else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_shading(cell, fill)
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(9)

    set_cell_borders(table)
    return table


# ---------------------------------------------------------------------------
# Markdown parser → docx
# ---------------------------------------------------------------------------

def parse_inline(text):
    """Return list of (text, bold, italic, code) tuples."""
    parts = []
    # Split on **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False, False))
        token = m.group(0)
        if token.startswith('**'):
            parts.append((token[2:-2], True, False, False))
        elif token.startswith('*'):
            parts.append((token[1:-1], False, True, False))
        else:
            parts.append((token[1:-1], False, False, True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False, False))
    return parts


def add_inline_para(doc, text, base_size=Pt(10), indent_left=None, shading=None, border_left_colour=None):
    para = doc.add_paragraph()
    if indent_left:
        ppr = para._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(indent_left))
        ppr.append(ind)
    if shading:
        set_para_shading(para, shading)
    if border_left_colour:
        add_para_border_left(para, border_left_colour, space=8, sz=16)
    for (chunk, bold, italic, code) in parse_inline(text):
        run = para.add_run(chunk)
        run.font.size = base_size
        run.font.bold = bold
        run.font.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GREY
    return para


def is_table_line(line):
    return line.strip().startswith('|') and line.strip().endswith('|')


def parse_table_block(lines, start):
    rows = []
    i = start
    while i < len(lines) and is_table_line(lines[i]):
        # Skip separator rows (---|---)
        if re.match(r'^\|[-| :]+\|$', lines[i].strip()):
            i += 1
            continue
        cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def render_md_to_docx(doc, md_text):
    lines = md_text.splitlines()
    i = 0
    in_code_block = False
    code_lines = []

    # State for Q&A detection
    prev_was_question = False

    while i < len(lines):
        line = lines[i]

        # ---- Code block ----
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, '\n'.join(code_lines))
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ---- Headings ----
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            add_heading(doc, text, level=min(level, 4))
            i += 1
            continue

        # ---- Horizontal rule ----
        if re.match(r'^---+$', line.strip()):
            doc.add_paragraph().paragraph_format.border_bottom = True
            i += 1
            continue

        # ---- Table ----
        if is_table_line(line):
            rows, i = parse_table_block(lines, i)
            if rows:
                add_table_from_rows(doc, rows[0], rows[1:])
                doc.add_paragraph()  # spacer
            continue

        # ---- Q&A detection ----
        # Lines starting with "> Q:" are questions; "> A:" are answers
        if line.startswith('> Q:') or re.match(r'^\*\*Q\d+', line):
            q_text = line.lstrip('> ').lstrip('*').rstrip('*').strip()
            add_question_para(doc, q_text)
            prev_was_question = True
            i += 1
            continue

        if prev_was_question and (line.startswith('> A:') or re.match(r'^\*\*A:', line)):
            a_text = line.lstrip('> ').lstrip('*A:').strip()
            add_answer_para(doc, a_text)
            prev_was_question = False
            i += 1
            continue

        prev_was_question = False

        # ---- Blockquote (generic) ----
        if line.startswith('> '):
            add_inline_para(doc, line[2:], indent_left=360,
                            shading=RGBColor(0xEB, 0xF4, 0xFB),
                            border_left_colour=ACCENT)
            i += 1
            continue

        # ---- Bullet list ----
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            indent = len(m.group(1))
            content = m.group(2)
            para = doc.add_paragraph(style='List Bullet')
            ppr = para._p.get_or_add_pPr()
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(360 + indent * 180))
            ind.set(qn('w:hanging'), '360')
            ppr.append(ind)
            for (chunk, bold, italic, code) in parse_inline(content):
                run = para.add_run(chunk)
                run.font.size = Pt(10)
                run.font.bold = bold
                run.font.italic = italic
                if code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            i += 1
            continue

        # ---- Numbered list ----
        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            content = m.group(1)
            para = doc.add_paragraph(style='List Number')
            for (chunk, bold, italic, code) in parse_inline(content):
                run = para.add_run(chunk)
                run.font.size = Pt(10)
                run.font.bold = bold
                run.font.italic = italic
                if code:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            i += 1
            continue

        # ---- Blank line ----
        if line.strip() == '':
            i += 1
            continue

        # ---- Normal paragraph ----
        add_inline_para(doc, line)
        i += 1


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def add_cover_page(doc):
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('Interview Prep Pack')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run('Backend Stage 4 — Insighta Labs+')
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = ACCENT

    # Spacer
    doc.add_paragraph()

    # Developer
    dev_para = doc.add_paragraph()
    dev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dev_run = dev_para.add_run('Developer: Kelechi Uba')
    dev_run.font.size = Pt(12)
    dev_run.font.color.rgb = DARK_GREY

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('2026-05-07')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = DARK_GREY

    # Page break
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    with open(MD_PATH, encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    add_cover_page(doc)

    # Strip the H1 title from the MD (already on cover) and render the rest
    body_md = re.sub(r'^# [^\n]+\n', '', md_text, count=1)
    # Skip the metadata block at top (lines starting with ** up to first ---)
    render_md_to_docx(doc, body_md)

    doc.save(DOCX_PATH)
    print(f'Saved: {DOCX_PATH}')


if __name__ == '__main__':
    main()
